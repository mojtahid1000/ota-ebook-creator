"""
Generate OTA-branded PRD document as DOCX and PDF.
OTA Ebook Creator - Product Requirements Document
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# OTA Brand Colors
OTA_ORANGE = RGBColor(0xFF, 0x6B, 0x35)
OTA_BLUE = RGBColor(0x1E, 0x3A, 0x8A)
OTA_TEAL = RGBColor(0x14, 0xB8, 0xA6)
OTA_DARK = RGBColor(0x1E, 0x29, 0x3B)
OTA_LIGHT_GRAY = RGBColor(0xF8, 0xFA, 0xFC)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(SCRIPT_DIR, "assets", "ota-logo-horizontal.png")
OUTPUT_DIR = os.path.join(SCRIPT_DIR)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex
    })
    shading.append(shading_elem)


def add_styled_table(doc, headers, rows, header_color="1E3A8A"):
    """Create OTA-branded table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(8.5)
            run.font.color.rgb = OTA_DARK
            if r % 2 == 1:
                set_cell_shading(cell, "F1F5F9")

    return table


def add_heading_styled(doc, text, level=1):
    """Add OTA-colored heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = OTA_BLUE if level == 1 else OTA_ORANGE
    return h


def add_para(doc, text, bold=False, color=None, size=10.5):
    """Add styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def create_prd_docx():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ===== COVER PAGE =====
    doc.add_paragraph()  # spacing

    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(2.5))

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI Ebook Creator")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = OTA_BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Product Requirements Document (PRD)")
    run.font.size = Pt(18)
    run.font.color.rgb = OTA_ORANGE

    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("ওটিএ ইবুক ক্রিয়েটর - এআই দিয়ে ইবুক তৈরি করুন")
    run.font.size = Pt(14)
    run.font.color.rgb = OTA_TEAL

    doc.add_paragraph()
    doc.add_paragraph()

    # Meta info
    meta_info = [
        ("Company:", "Online Tech Academy (OTA)"),
        ("Founder:", "Mentor Mojtahidul Islam (@mentormojtahid)"),
        ("Website:", "www.onlinetechacademy.com"),
        ("Version:", "1.0"),
        ("Date:", "March 19, 2026"),
        ("Document Type:", "Product Requirements Document"),
    ]
    for label, value in meta_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label + " ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = OTA_BLUE
        run = p.add_run(value)
        run.font.size = Pt(11)
        run.font.color.rgb = OTA_DARK

    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    add_heading_styled(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Product Overview & Vision",
        "2. OTA Brand Guide",
        "3. System Architecture (12-Agent Pipeline)",
        "4. Agent Specifications (All 12 Agents)",
        "5. Agent Forward Data Chain",
        "6. User Flow (Step-by-Step)",
        "7. ContentWriter - Topic-by-Topic Flow",
        "8. Writing Styles & Rewrite Angles",
        "9. Font Selection & Page Design Options",
        "10. Responsive Reading Experience",
        "11. Token/Credit Management & Auto-Resume",
        "12. Database Schema",
        "13. Financial Model (Costs, Pricing, Profit/Loss)",
        "14. Knowledge Files & AI Data Specs",
        "15. Build Phases & Timeline",
        "16. Resources & Accounts Required",
        "17. Verification Checklist",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        run = p.runs[0]
        run.font.size = Pt(11)
        run.font.color.rgb = OTA_DARK

    doc.add_page_break()

    # ===== 1. PRODUCT OVERVIEW =====
    add_heading_styled(doc, "1. Product Overview & Vision", level=1)

    add_para(doc, "Product Name: OTA Ebook Creator (ওটিএ ইবুক ক্রিয়েটর)", bold=True, color=OTA_ORANGE, size=12)
    add_para(doc,
        "A step-by-step AI-powered ebook creation platform where users select a niche, "
        "define their target audience, and let 12 specialized AI agents write a complete "
        "problem-solution ebook - topic by topic, with full control over content, design, and export."
    )
    add_para(doc,
        "The platform uses Claude Agent SDK (Python) with each step handled by a specialist agent. "
        "Users confirm each step before the next agent runs. The ContentWriter agent writes one topic "
        "at a time, allowing users to choose writing styles, request rewrites with different angles, "
        "and confirm page designs before proceeding."
    )

    add_heading_styled(doc, "Key Features", level=2)
    features = [
        "12 specialized AI agents (each optimized for its task)",
        "Topic-by-topic writing with style selection and rewrite options",
        "7 Bangla font choices with customizable header/footer/page design",
        "Responsive ebook preview (laptop + mobile comfortable reading)",
        "AI-generated book covers with Bangla text overlay",
        "Export to PDF, DOCX, and Google Docs with full styling",
        "Smart token management with auto-pause and auto-resume",
        "Book title + tagline generator (8-10 options)",
        "Freemium model: Free / Pro (BDT 999) / Unlimited (BDT 2,499)",
        "SSLCommerz + bKash payment integration (Bangladesh local)",
        "All UI in Bangla with English technical terms",
    ]
    for f in features:
        p = doc.add_paragraph(f, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_page_break()

    # ===== 2. BRAND GUIDE =====
    add_heading_styled(doc, "2. OTA Brand Guide", level=1)

    add_heading_styled(doc, "Brand Colors", level=2)
    add_styled_table(doc,
        ["Color", "Hex", "RGB", "Usage"],
        [
            ["OTA Orange (Primary)", "#FF6B35", "rgb(255, 107, 53)", "Headers, CTAs, primary buttons"],
            ["OTA Deep Blue (Secondary)", "#1E3A8A", "rgb(30, 58, 138)", "Subheadings, backgrounds, dividers"],
            ["OTA Teal (Accent)", "#14B8A6", "rgb(20, 184, 166)", "Progress bars, success states, accents"],
            ["White", "#FFFFFF", "rgb(255, 255, 255)", "Backgrounds, text on dark"],
            ["Light Gray", "#F8FAFC", "rgb(248, 250, 252)", "Card backgrounds, subtle borders"],
            ["Dark Text", "#1E293B", "rgb(30, 41, 59)", "Body text"],
        ]
    )

    doc.add_paragraph()
    add_heading_styled(doc, "Typography", level=2)
    add_styled_table(doc,
        ["Element", "Font", "Weight", "Size"],
        [
            ["Web Headings", "Hind Siliguri", "Bold (700)", "24-32px"],
            ["Web Sub-headings", "Hind Siliguri", "SemiBold (600)", "18-22px"],
            ["Web Body", "Hind Siliguri", "Regular (400)", "16px"],
            ["Ebook Headings", "Calibri Bold / Hind Siliguri Bold", "700", "18-24pt"],
            ["Ebook Body", "Calibri / Hind Siliguri", "400", "11-12pt"],
        ]
    )

    doc.add_paragraph()
    add_heading_styled(doc, "Brand Voice", level=2)
    add_para(doc, "Tone: Professional, educational, and supportive", bold=True)
    add_para(doc, "Style: Clear, direct, and instructive - 'Knowledgeable Mentor'")
    add_para(doc, "Language: Bangla with English technical terms in parentheses")

    doc.add_paragraph()
    add_heading_styled(doc, "Logo Assets", level=2)
    logos = [
        "ota-logo-horizontal.png - Header/cover use",
        "ota-logo-square.png - Square variant",
        "ota-logo-white.png - White version for dark backgrounds",
    ]
    for l in logos:
        doc.add_paragraph(l, style='List Bullet')

    doc.add_page_break()

    # ===== 3. ARCHITECTURE =====
    add_heading_styled(doc, "3. System Architecture", level=1)

    add_para(doc, "Frontend: Next.js 14+ (App Router) + Tailwind CSS + TypeScript", bold=True, size=11)
    add_para(doc, "Backend: Python FastAPI + Claude Agent SDK", bold=True, size=11)
    add_para(doc, "Database: Supabase (PostgreSQL + Auth + Storage)", bold=True, size=11)
    add_para(doc, "AI: Claude API (Anthropic) + OpenAI GPT (user chooses)", bold=True, size=11)
    add_para(doc, "Payments: SSLCommerz + bKash (Bangladesh local)", bold=True, size=11)
    add_para(doc, "Image Gen: DALL-E 3 + fal.ai + inference.sh", bold=True, size=11)

    doc.add_paragraph()
    add_para(doc,
        "Architecture Flow: Next.js frontend communicates with Python FastAPI backend via REST + SSE streaming. "
        "The backend orchestrates 12 specialized AI agents using Claude Agent SDK. Each agent receives input from "
        "the previous agent, processes it, and forwards structured output to the next. User confirms each step "
        "via the frontend before the next agent is triggered.",
        size=10
    )

    doc.add_page_break()

    # ===== 4. AGENT SPECIFICATIONS =====
    add_heading_styled(doc, "4. Agent Specifications (All 12 Agents)", level=1)

    agents = [
        {
            "num": 1, "name": "NicheNavigator", "model": "claude-haiku-4-5",
            "task": "Niche & sub-niche selection with market demand data",
            "why": "Simple selection task - fast, low cost",
            "rules": [
                "Respond in Bangla with English terms in parentheses",
                "Sub-niches must be specific enough for a focused ebook",
                "Include market demand indicator (High/Medium/Low)",
                "Maximum 40 sub-niches per category",
            ],
            "output": "main_niche, sub_niche, niche_description, market_demand, target_keywords[]"
        },
        {
            "num": 2, "name": "AvatarArchitect", "model": "claude-sonnet-4-6",
            "task": "Build detailed reader persona with demographics + psychographics",
            "why": "Needs reasoning for persona creation",
            "rules": [
                "Generate age ranges relevant to sub-niche (not generic)",
                "Pain points must be specific and emotionally resonant",
                "Include surface-level and deep emotional pain points",
                "Generate 8-12 specific pain points per avatar",
                "All labels in Bangla",
            ],
            "output": "avatar: { age_range, gender, income, education, pain_points[], goals[], emotional_triggers[] }"
        },
        {
            "num": 3, "name": "ProblemDetective", "model": "claude-sonnet-4-6",
            "task": "Discover & rank 10 specific problems based on avatar + niche",
            "why": "Analytical problem discovery requires reasoning",
            "rules": [
                "Generate exactly 10 problems ranked by ebook potential",
                "Problems must be specific (not generic like 'lose weight')",
                "Include urgency score (1-10) and emotional weight (1-10)",
                "Include 'why this is a great ebook topic' for each",
                "Bangla with English medical/technical terms",
            ],
            "output": "selected_problem: { title, description, urgency, emotional_weight, why_great_topic }"
        },
        {
            "num": 4, "name": "SolutionStrategist", "model": "claude-sonnet-4-6",
            "task": "Propose categorized solution approaches with pros/cons",
            "why": "Domain expertise reasoning needed",
            "rules": [
                "Generate 5-8 solution categories",
                "Each has: name, description, pros, cons, audience fit score",
                "At least one natural/holistic and one conventional option",
                "User can select 1-3 approaches",
                "Flag approaches needing medical disclaimers",
            ],
            "output": "selected_solutions[]: { name, description, pros, cons, estimated_pages }"
        },
        {
            "num": 5, "name": "ResearchAnalyst", "model": "claude-opus-4-6",
            "task": "Deep research: statistics, expert quotes, case studies",
            "why": "Deep reasoning, accuracy critical - needs best model",
            "rules": [
                "Research must be factual and verifiable",
                "Include 10-15 key statistics with sources",
                "Include 3-5 expert quotes or references",
                "Include 2-3 case studies",
                "Add Bangladesh/South Asia context",
                "Include Research Confidence Score",
            ],
            "output": "research: { summary, statistics[], quotes[], case_studies[], confidence_score }"
        },
        {
            "num": 6, "name": "BookNameCreator", "model": "claude-sonnet-4-6",
            "task": "Generate 8-10 book title + tagline options with different formulas",
            "why": "Creative + analytical reasoning for compelling titles",
            "rules": [
                "8-10 title + tagline pairs using different formulas",
                "Formulas: How-To, Number, Question, Promise, Story, Contrarian, Metaphor, One-Word",
                "Tagline max 12 words",
                "Titles specific to the problem/solution",
                "Bangla with English version in brackets",
                "Include Title Strength Score (1-10)",
            ],
            "output": "book_title, book_subtitle, tagline, title_style"
        },
        {
            "num": 7, "name": "OutlineArchitect", "model": "claude-sonnet-4-6",
            "task": "Design complete ebook structure with chapters, topics, page estimates",
            "why": "Structural/organizational reasoning",
            "rules": [
                "10-15 chapters with logical flow",
                "Each chapter: title, 3-5 sub-topics, estimated pages, content types",
                "First = hook/intro, last = action plan/conclusion",
                "Include 2+ chapters with exercises, 1 with checklist",
                "Total 80-150 pages",
                "Chapter titles must be compelling",
            ],
            "output": "outline: { chapters[]: { title, topics[], estimated_pages, content_types[] } }"
        },
        {
            "num": 8, "name": "ContentWriter", "model": "claude-opus-4-6",
            "task": "Write ONE topic at a time with style options, confirm each before next",
            "why": "Highest quality writing - needs best model",
            "rules": [
                "Write ONE topic per invocation",
                "Ask user for writing style preference first",
                "Include: headline, intro hook, subheadings, quotes, stats, tips",
                "2000-4000 words per topic depending on estimated pages",
                "Use markdown formatting",
                "Include 'Pro Tips' callout boxes",
                "User confirms each topic before writing next",
                "Support rewrite with different angle",
                "NO plagiarism - all original content",
            ],
            "output": "chapters_content[]: { title, content_markdown, word_count, writing_style }"
        },
        {
            "num": 9, "name": "EditorReviewer", "model": "claude-sonnet-4-6",
            "task": "Quality review: consistency, flow, scoring, revision flags",
            "why": "Careful analytical review",
            "rules": [
                "Check chapter-to-chapter flow, terminology, tone consistency",
                "Verify all outlined sub-topics are covered",
                "Score: quality, readability, engagement, actionability (1-10 each)",
                "Flag chapters needing revision with specific feedback",
                "Generate book description for back cover",
            ],
            "output": "review: { quality_score, readability_score, revision_flags[], book_description }"
        },
        {
            "num": 10, "name": "ExportMaster", "model": "claude-haiku-4-5",
            "task": "Format & export: PDF/DOCX/Google Docs with design options",
            "why": "Fast coordination, formatting rules",
            "rules": [
                "Present font/header/footer/decoration options to user",
                "Generate complete TOC with page mapping",
                "Include: title page, copyright, TOC, chapters, about author",
                "Apply consistent design across all pages",
                "PDF: A4 or 6x9 standard ebook size",
            ],
            "output": "export: { formats[], file_urls{}, design_settings }"
        },
        {
            "num": 11, "name": "CoverDesigner", "model": "claude-sonnet-4-6 + DALL-E 3 / fal.ai",
            "task": "AI cover art generation + Bangla text overlay",
            "why": "Creative prompt engineering + image gen",
            "rules": [
                "Generate 3-4 front cover variations",
                "AI generates background art ONLY (no text in image)",
                "Bangla text overlay added programmatically (Pillow)",
                "Color schemes match niche mood",
                "Back cover: description, author bio, press, website",
                "Cover dimensions: 1600x2400px",
                "Supports DALL-E 3, fal.ai, inference.sh",
            ],
            "output": "covers: { front_options[], back, selected_front_index }"
        },
        {
            "num": 12, "name": "DeliveryManager", "model": "claude-haiku-4-5",
            "task": "Final assembly, verify all files, create download package",
            "why": "Fast, simple coordination",
            "rules": [
                "Verify all components exist",
                "Create final download package",
                "Generate Creation Summary",
                "Store all files in Supabase Storage",
                "Update project status to completed",
                "Show congratulations in Bangla",
            ],
            "output": "downloads: { pdf_url, docx_url, gdocs_url, front_cover_url, back_cover_url }"
        },
    ]

    for agent in agents:
        add_heading_styled(doc, f"Agent {agent['num']}: {agent['name']}", level=2)
        add_para(doc, f"Model: {agent['model']}", bold=True, color=OTA_TEAL, size=10)
        add_para(doc, f"Why this model: {agent['why']}", size=9.5)
        add_para(doc, f"Task: {agent['task']}", bold=True, size=10)

        add_para(doc, "Rules:", bold=True, color=OTA_ORANGE, size=10)
        for rule in agent['rules']:
            p = doc.add_paragraph(rule, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(9)

        add_para(doc, f"Output: {agent['output']}", bold=True, size=9.5, color=OTA_BLUE)
        doc.add_paragraph()

    doc.add_page_break()

    # ===== 5. FORWARD CHAIN =====
    add_heading_styled(doc, "5. Agent Forward Data Chain", level=1)

    chain = [
        ("Agent 1 -> 2", "main_niche, sub_niche, niche_description, market_demand, target_keywords[]"),
        ("Agent 2 -> 3", "+ avatar { age, gender, income, education, pain_points[], goals[], emotional_triggers[] }"),
        ("Agent 3 -> 4", "+ selected_problem { title, description, urgency, emotional_weight }"),
        ("Agent 4 -> 5", "+ selected_solutions[] { name, description, pros, cons, estimated_pages }"),
        ("Agent 5 -> 6", "+ research { summary, statistics[], quotes[], case_studies[], confidence_score }"),
        ("Agent 6 -> 7", "+ book_title, book_subtitle, tagline"),
        ("Agent 7 -> 8", "+ outline { chapters[] with topics[], estimated_pages, content_types[] }"),
        ("Agent 8 -> 9", "+ chapters_content[] { content_markdown, word_count } (all confirmed topics)"),
        ("Agent 9 -> 10", "+ review { quality_scores, revision_flags[], book_description }"),
        ("Agent 10 -> 11", "+ export { format, file_url, design_settings }"),
        ("Agent 11 -> 12", "+ covers { front_url, back_url, selected_style }"),
        ("Agent 12 -> User", "FINAL: all download links + creation summary"),
    ]
    add_styled_table(doc,
        ["From -> To", "Data Forwarded"],
        chain,
        header_color="FF6B35"
    )

    doc.add_page_break()

    # ===== 7. CONTENT WRITER FLOW =====
    add_heading_styled(doc, "6. ContentWriter - Topic-by-Topic Flow", level=1)

    add_para(doc,
        "The ContentWriter (Agent 8) is the most complex agent. It writes ONE topic at a time, "
        "asks for user's writing style preference, and supports rewrite with different angles. "
        "Every topic's page design must be confirmed before proceeding.",
        size=10
    )

    add_heading_styled(doc, "Writing Flow Per Topic", level=2)
    steps = [
        "1. Ask user: 'What writing style do you prefer?' (show 7 options)",
        "2. Write topic content with selected style",
        "3. Show topic with page design preview",
        "4. User choices:",
        "   a) Confirm -> Save, move to next topic",
        "   b) Rewrite (same style) -> Regenerate",
        "   c) Rewrite (different angle) -> Show angle options, then rewrite",
        "   d) Edit manually -> User edits in-place",
        "   e) Go back -> Return to previous topic",
        "5. After confirm -> automatically start next topic",
    ]
    for s in steps:
        p = doc.add_paragraph(s)
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_paragraph()
    add_heading_styled(doc, "7 Writing Styles", level=2)
    add_styled_table(doc,
        ["Style", "Bangla Label", "Description"],
        [
            ["Storytelling", "গল্পের ধরনে", "Narrative with relatable stories and examples"],
            ["Step-by-Step Guide", "ধাপে ধাপে গাইড", "Numbered instructions, clear action items"],
            ["Conversational", "কথোপকথনের ধরনে", "Casual, like talking to a friend"],
            ["Academic", "একাডেমিক", "Research-backed, formal with citations"],
            ["Motivational", "অনুপ্রেরণামূলক", "Encouraging, empowering, with quotes"],
            ["Q&A Format", "প্রশ্ন-উত্তর", "Written as questions and detailed answers"],
            ["Case Study", "কেস স্টাডি", "Real examples and analysis-focused"],
        ]
    )

    doc.add_paragraph()
    add_heading_styled(doc, "7 Rewrite Angles", level=2)
    add_styled_table(doc,
        ["Angle", "Description"],
        [
            ["More practical", "Focus on actionable tips, less theory"],
            ["More emotional", "Connect with reader's feelings and struggles"],
            ["More data-driven", "Add more statistics and research"],
            ["Simpler language", "Make it easier to understand"],
            ["More detailed", "Expand with deeper explanations"],
            ["Different example", "Use completely different stories/examples"],
            ["Shorter version", "Condense to key points only"],
        ]
    )

    doc.add_page_break()

    # ===== FONT & DESIGN =====
    add_heading_styled(doc, "7. Font Selection & Page Design Options", level=1)

    add_heading_styled(doc, "Bangla Font Options (User Selects)", level=2)
    add_styled_table(doc,
        ["Font", "Style", "Best For"],
        [
            ["Hind Siliguri", "Clean, modern sans-serif", "General ebooks, guides"],
            ["Noto Sans Bengali", "Google standard, wide coverage", "Technical ebooks"],
            ["Kalpurush", "Traditional, elegant", "Literary/emotional topics"],
            ["SolaimanLipi", "Classic, widely used", "Formal/academic ebooks"],
            ["Adarsha Lipi", "Decorative, artistic", "Creative/lifestyle ebooks"],
            ["Mukti", "Simple, readable", "Long-form reading"],
            ["Bangla MN", "Apple system, clean", "Apple device readers"],
        ]
    )

    doc.add_paragraph()
    add_heading_styled(doc, "Header Styles (5 Options)", level=2)
    headers_list = [
        "1. Minimal: Book title (left) | Chapter name (right)",
        "2. Branded: OTA logo (left) | Chapter name (center) | Page # (right)",
        "3. Bold Line: Full-width colored line under header text",
        "4. Decorative: Ornamental border with centered title",
        "5. None: No header",
    ]
    for h in headers_list:
        doc.add_paragraph(h, style='List Bullet')

    add_heading_styled(doc, "Footer Styles (5 Options)", level=2)
    footers_list = [
        "1. Standard: Author name (left) | Page # (center) | Website (right)",
        "2. Minimal: Just page number centered",
        "3. Branded: OTA logo (left) | Page # (center) | Press name (right)",
        "4. Line Above: Thin line above footer content",
        "5. Copyright: (c) Author | Press | Page #",
    ]
    for f in footers_list:
        doc.add_paragraph(f, style='List Bullet')

    add_heading_styled(doc, "Page Decoration Options", level=2)
    add_styled_table(doc,
        ["Element", "Options"],
        [
            ["Chapter Title Pages", "Full page decorative / Half page / Simple heading"],
            ["Quote Boxes", "Rounded border / Left accent bar / Shadowed / Italic only"],
            ["Tip/Pro Boxes", "Orange gradient / Blue sidebar / Icon-prefixed / Dashed border"],
            ["Stat Callouts", "Large number / Chart-style / Highlighted background"],
            ["Page Borders", "None / Thin line / Decorative corner / Full frame"],
            ["Margins", "Narrow (0.5in) / Standard (0.75in) / Wide (1in) / Extra wide (1.25in)"],
            ["Line Spacing", "Tight (1.2) / Normal (1.5) / Relaxed (1.8) / Double (2.0)"],
        ]
    )

    doc.add_page_break()

    # ===== RESPONSIVE =====
    add_heading_styled(doc, "8. Responsive Reading Experience", level=1)

    add_heading_styled(doc, "Reading Comfort Features", level=2)
    comfort = [
        "Font size slider (14px - 24px)",
        "Dark mode / Light mode / Sepia mode",
        "Reading progress bar at top",
        "Estimated reading time per chapter",
        "Swipe navigation on mobile",
        "Text max-width: 700px for optimal reading (laptop)",
        "Proper line-height: 1.6-1.8 for Bangla text",
        "Pinch to zoom on mobile",
        "Bookmark any section",
    ]
    for c in comfort:
        doc.add_paragraph(c, style='List Bullet')

    add_heading_styled(doc, "Laptop View (1024px+)", level=2)
    add_para(doc, "Left sidebar: Table of Contents with chapter status icons. "
        "Main area: centered ebook page view with 700px max text width. "
        "Design and Export buttons in sidebar.", size=10)

    add_heading_styled(doc, "Mobile View (375px)", level=2)
    add_para(doc, "Hamburger menu for TOC. Full-width content with larger font (18px). "
        "Touch scroll with swipe for next/prev chapter. Bottom nav bar.", size=10)

    doc.add_page_break()

    # ===== TOKEN MANAGEMENT =====
    add_heading_styled(doc, "9. Token/Credit Management & Auto-Resume", level=1)

    add_para(doc,
        "Smart token management ensures no work is lost if API credits run out mid-generation. "
        "Each topic completion is a checkpoint. The system can pause and resume from the exact stopping point.",
        size=10
    )

    add_heading_styled(doc, "Token Budget Per Agent", level=2)
    add_styled_table(doc,
        ["Agent", "Est. Tokens/Call", "Priority"],
        [
            ["NicheNavigator", "~500", "Low"],
            ["AvatarArchitect", "~800", "Low"],
            ["ProblemDetective", "~1,200", "Medium"],
            ["SolutionStrategist", "~1,000", "Medium"],
            ["ResearchAnalyst", "~3,000", "High"],
            ["BookNameCreator", "~800", "Low"],
            ["OutlineArchitect", "~2,000", "Medium"],
            ["ContentWriter", "~4,000/topic", "Critical"],
            ["EditorReviewer", "~2,500", "Medium"],
            ["ExportMaster", "~500", "Low"],
            ["CoverDesigner", "~1,000 + DALL-E", "High"],
            ["DeliveryManager", "~300", "Low"],
        ]
    )

    add_para(doc, "Total per ebook (~36 topics): ~157,000 tokens", bold=True, color=OTA_ORANGE)

    add_heading_styled(doc, "Auto-Resume Strategy", level=2)
    add_para(doc,
        "Before each AI call: check remaining credits -> estimate tokens needed -> if sufficient, proceed -> "
        "if insufficient, PAUSE and save state to DB. On resume: read saved context, reconstruct agent state, "
        "continue from exact stop point. The ContentWriter reads a summary of all previous topics and continues "
        "writing from where it stopped.",
        size=10
    )

    doc.add_page_break()

    # ===== FINANCIAL MODEL =====
    add_heading_styled(doc, "10. Financial Model", level=1)

    add_heading_styled(doc, "A. Cost Per Ebook Generated", level=2)
    add_styled_table(doc,
        ["Component", "Cost (USD)", "Cost (BDT)"],
        [
            ["ContentWriter (Opus, ~144K tokens)", "$1.20-1.80", "BDT 132-198"],
            ["ResearchAnalyst (Opus, ~3K tokens)", "$0.05", "BDT 5.5"],
            ["Other 10 agents (Sonnet/Haiku)", "$0.15-0.30", "BDT 16-33"],
            ["DALL-E 3 covers (4 images)", "$0.16", "BDT 17.6"],
            ["Storage (PDF/DOCX/images)", "$0.01", "BDT 1.1"],
            ["TOTAL PER EBOOK", "$1.57-2.32", "BDT 173-255"],
        ],
        header_color="FF6B35"
    )

    doc.add_paragraph()
    add_heading_styled(doc, "B. Pricing Plans", level=2)
    add_styled_table(doc,
        ["Plan", "Price/Month (BDT)", "Price/Month (USD)", "Includes"],
        [
            ["Free", "0 (Free)", "$0", "1 ebook/month, PDF only, basic fonts"],
            ["Pro", "999", "$9", "10 ebooks/month, all formats, AI covers, all fonts"],
            ["Unlimited", "2,499", "$22", "Unlimited ebooks, all features, custom branding"],
        ]
    )

    doc.add_paragraph()
    add_heading_styled(doc, "C. Revenue Projections", level=2)
    add_styled_table(doc,
        ["Scenario", "Free Users", "Pro Users", "Unlimited", "Revenue (BDT)", "Costs (BDT)", "Profit (BDT)"],
        [
            ["Month 1 (Launch)", "200", "10", "2", "14,988", "25,000", "-10,012"],
            ["Month 3", "500", "40", "8", "59,952", "35,000", "+24,952"],
            ["Month 6", "1,000", "100", "25", "162,350", "55,000", "+107,350"],
            ["Month 12", "2,000", "300", "80", "499,620", "100,000", "+399,620"],
            ["Year 2", "5,000", "800", "200", "1,298,800", "200,000", "+1,098,800"],
        ]
    )

    doc.add_paragraph()
    add_heading_styled(doc, "D. Break-Even Analysis", level=2)
    add_para(doc, "Break-Even ROAS = 1 / Gross Margin % = 1 / 0.70 = 1.43x", bold=True)
    add_para(doc, "Break-Even Point = Fixed Costs / Contribution Margin = 25,000 / 799 = 32 Pro subscribers", bold=True, color=OTA_TEAL)

    doc.add_paragraph()
    add_heading_styled(doc, "E. Unit Economics", level=2)
    add_styled_table(doc,
        ["Metric", "Value"],
        [
            ["Cost per ebook", "BDT 173-255 (~$1.57-2.32)"],
            ["Revenue per Pro ebook", "BDT 99.9 (999/10)"],
            ["Gross margin (Pro)", "~60-75%"],
            ["Customer Acquisition Cost (target)", "BDT 500-1,000"],
            ["Lifetime Value (6-month avg)", "BDT 5,994"],
            ["LTV:CAC Ratio", "6:1 - 12:1"],
        ]
    )

    doc.add_page_break()

    # ===== PLATFORM COSTS =====
    add_heading_styled(doc, "F. Platform Operating Costs (Monthly)", level=2)
    add_styled_table(doc,
        ["Cost Item", "Monthly (USD)", "Monthly (BDT)", "Notes"],
        [
            ["Claude Opus (Writer+Research)", "$50-200", "BDT 5,500-22,000", "Best model for quality"],
            ["Claude Sonnet (6 agents)", "$20-80", "BDT 2,200-8,800", "Mid-tier reasoning"],
            ["Claude Haiku (3 agents)", "$5-15", "BDT 550-1,650", "Fast, cheap"],
            ["OpenAI GPT-4o (alt option)", "$30-100", "BDT 3,300-11,000", "If user selects OpenAI"],
            ["DALL-E 3 (covers)", "$10-40", "BDT 1,100-4,400", "~$0.04/image"],
            ["Supabase Pro", "$25", "BDT 2,750", "DB + Auth + Storage"],
            ["Vercel Pro", "$20", "BDT 2,200", "Frontend hosting"],
            ["Python Backend", "$20-50", "BDT 2,200-5,500", "Railway/Render"],
            ["Domain + SSL", "$2", "BDT 220", "Annual spread"],
            ["TOTAL (100 users)", "$182-512", "BDT 20,000-56,000", ""],
            ["TOTAL (1000 users)", "$500-1,500", "BDT 55,000-165,000", ""],
        ]
    )

    doc.add_page_break()

    # ===== KNOWLEDGE FILES =====
    add_heading_styled(doc, "11. Knowledge Files & AI Data Specs", level=1)

    add_heading_styled(doc, "Pre-Built Knowledge Files (Created Before Launch)", level=2)
    add_styled_table(doc,
        ["File", "Path", "Content", "Size"],
        [
            ["health_niches_bn.json", "knowledge/niches/", "40+ health sub-niches in Bangla", "~15KB"],
            ["wealth_niches_bn.json", "knowledge/niches/", "35+ wealth sub-niches", "~12KB"],
            ["relationship_niches_bn.json", "knowledge/niches/", "30+ relationship sub-niches", "~10KB"],
            ["demographics_bn.json", "knowledge/avatars/", "Age, gender, income, education in Bangla", "~5KB"],
            ["pain_points_templates.json", "knowledge/avatars/", "50+ pain points per niche", "~20KB"],
            ["title_formulas.json", "knowledge/titles/", "10 title formulas with examples", "~8KB"],
            ["chapter_templates.json", "knowledge/writing/", "Content structure per style", "~15KB"],
            ["bangla_style_guide.md", "knowledge/writing/", "Bangla writing rules, 70/30 ratio", "~5KB"],
            ["dalle_prompt_templates.json", "knowledge/covers/", "Prompt templates per niche", "~10KB"],
            ["color_schemes.json", "knowledge/covers/", "20+ palettes by niche mood", "~5KB"],
            ["font_specimens.json", "knowledge/design/", "Bangla font metadata", "~8KB"],
            ["page_design_presets.json", "knowledge/design/", "Header/footer/decoration configs", "~12KB"],
        ]
    )

    doc.add_page_break()

    # ===== BUILD PHASES =====
    add_heading_styled(doc, "12. Build Phases & Timeline", level=1)

    phases = [
        ("Phase 1: Foundation", "Days 1-4",
         "Next.js + Python FastAPI + Supabase + Auth + Dashboard + Knowledge files"),
        ("Phase 2: Agents 1-4 + UI", "Days 5-9",
         "NicheNavigator, AvatarArchitect, ProblemDetective, SolutionStrategist + Steps 1-5 UI"),
        ("Phase 3: Agents 5-7 + UI", "Days 10-14",
         "ResearchAnalyst (Opus), BookNameCreator, OutlineArchitect + Steps 6-8 UI"),
        ("Phase 4: Agent 8 ContentWriter", "Days 15-21",
         "Topic-by-topic writing, style selection, rewrite angles, page design, token management"),
        ("Phase 5: Agents 9-10 Export", "Days 22-26",
         "EditorReviewer, ExportMaster, PDF/DOCX/Google Docs, font selection, design options"),
        ("Phase 6: Agents 11-12 + Payments", "Days 27-32",
         "CoverDesigner (DALL-E + fal.ai), DeliveryManager, SSLCommerz/bKash integration"),
        ("Phase 7: Polish + Launch", "Days 33-37",
         "Mobile responsive, error handling, auto-resume testing, landing page, final QA"),
    ]
    add_styled_table(doc,
        ["Phase", "Timeline", "Deliverables"],
        [(p[0], p[1], p[2]) for p in phases]
    )

    doc.add_page_break()

    # ===== RESOURCES =====
    add_heading_styled(doc, "13. Resources & Accounts Required", level=1)

    add_heading_styled(doc, "Required (Must Have)", level=2)
    add_styled_table(doc,
        ["Resource", "Purpose", "How to Get"],
        [
            ["Supabase Account", "Database, Auth, Storage", "https://supabase.com"],
            ["Anthropic API Key", "Claude AI (all agents)", "https://console.anthropic.com"],
            ["OpenAI API Key", "GPT-4o + DALL-E 3", "https://platform.openai.com"],
            ["SSLCommerz Account", "Primary payment gateway (BD)", "https://sslcommerz.com"],
            ["bKash Merchant", "Mobile payment (BD)", "https://developer.bkash.com"],
            ["Node.js 18+", "Next.js frontend", "https://nodejs.org"],
            ["Python 3.11+", "FastAPI + Agent SDK", "https://python.org"],
        ]
    )

    doc.add_paragraph()
    add_heading_styled(doc, "Optional (Add Later)", level=2)
    add_styled_table(doc,
        ["Resource", "Purpose", "How to Get"],
        [
            ["fal.ai API Key", "Alternative image gen", "https://fal.ai"],
            ["inference.sh API Key", "Alternative AI models", "https://inference.sh"],
            ["Google Cloud Project", "Docs export API", "https://console.cloud.google.com"],
            ["Vercel Account", "Frontend deploy", "https://vercel.com"],
            ["Railway/Render", "Backend deploy", "https://railway.app"],
        ]
    )

    doc.add_page_break()

    # ===== VERIFICATION =====
    add_heading_styled(doc, "14. Verification Checklist", level=1)

    checks = [
        "All 12 agents run and chain correctly",
        "Each agent output saved to DB before forwarding",
        "User can go BACK at any step",
        "ContentWriter writes topic-by-topic with confirm flow",
        "Rewrite with different angle works",
        "7 writing style selections work",
        "Token manager pauses when credits low",
        "Auto-resume reads context and continues correctly",
        "PDF exports with correct Bangla font, TOC, headers, footers, page numbers",
        "Responsive reading view on laptop (700px) and mobile (full-width)",
        "Font selection applies correctly to exported ebook",
        "Header/footer design options render in preview and export",
        "Book title generator shows 8-10 options with taglines",
        "Cover generator produces 3-4 DALL-E options",
        "SSLCommerz/bKash payments work (Free/Pro/Unlimited)",
        "All UI labels in Bangla",
        "OTA brand colors applied throughout",
        "Page design confirmation per topic",
        "Dark mode / Light mode / Sepia mode in preview",
    ]
    for check in checks:
        p = doc.add_paragraph(check, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)

    # ===== FOOTER =====
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Online Tech Academy | www.onlinetechacademy.com | @mentormojtahid")
    run.font.size = Pt(9)
    run.font.color.rgb = OTA_BLUE
    run.italic = True

    # Save
    docx_path = os.path.join(OUTPUT_DIR, "OTA_Ebook_Creator_PRD.docx")
    doc.save(docx_path)
    print(f"DOCX saved: {docx_path}")
    return docx_path


def create_prd_pdf(docx_path):
    """Generate a styled PDF version of the PRD."""
    from fpdf import FPDF

    class OTAPDF(FPDF):
        def header(self):
            self.set_font('Helvetica', 'B', 8)
            self.set_text_color(30, 58, 138)
            self.cell(0, 8, 'OTA Ebook Creator - PRD v1.0', 0, 0, 'L')
            self.cell(0, 8, 'Online Tech Academy', 0, 1, 'R')
            self.set_draw_color(255, 107, 53)
            self.set_line_width(0.5)
            self.line(10, 14, 200, 14)
            self.ln(4)

        def footer(self):
            self.set_y(-15)
            self.set_font('Helvetica', 'I', 8)
            self.set_text_color(30, 58, 138)
            self.cell(0, 10, f'www.onlinetechacademy.com | Page {self.page_no()}/{{nb}}', 0, 0, 'C')

    pdf = OTAPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Cover page
    pdf.add_page()
    pdf.set_font('Helvetica', '', 12)

    if os.path.exists(LOGO_PATH):
        pdf.image(LOGO_PATH, x=65, y=30, w=80)

    pdf.ln(70)
    pdf.set_font('Helvetica', 'B', 32)
    pdf.set_text_color(30, 58, 138)
    pdf.cell(0, 15, 'AI Ebook Creator', 0, 1, 'C')

    pdf.set_font('Helvetica', '', 16)
    pdf.set_text_color(255, 107, 53)
    pdf.cell(0, 10, 'Product Requirements Document', 0, 1, 'C')

    pdf.ln(10)
    pdf.set_font('Helvetica', '', 11)
    pdf.set_text_color(30, 41, 59)
    pdf.cell(0, 8, 'Company: Online Tech Academy (OTA)', 0, 1, 'C')
    pdf.cell(0, 8, 'Founder: Mentor Mojtahidul Islam', 0, 1, 'C')
    pdf.cell(0, 8, 'Website: www.onlinetechacademy.com', 0, 1, 'C')
    pdf.cell(0, 8, 'Version: 1.0 | Date: March 19, 2026', 0, 1, 'C')

    # Content pages
    sections = [
        ("12-Agent Pipeline Architecture", [
            "The system uses 12 specialized AI agents, each optimized for its specific task.",
            "Each agent forwards structured output to the next agent in the chain.",
            "User confirms each step before the next agent runs.",
            "",
            "Agent 1: NicheNavigator (Haiku) - Niche & sub-niche selection",
            "Agent 2: AvatarArchitect (Sonnet) - Target reader persona",
            "Agent 3: ProblemDetective (Sonnet) - Problem discovery & ranking",
            "Agent 4: SolutionStrategist (Sonnet) - Solution approaches",
            "Agent 5: ResearchAnalyst (Opus) - Deep research & statistics",
            "Agent 6: BookNameCreator (Sonnet) - Title + tagline generation",
            "Agent 7: OutlineArchitect (Sonnet) - Ebook structure & TOC",
            "Agent 8: ContentWriter (Opus) - Topic-by-topic writing",
            "Agent 9: EditorReviewer (Sonnet) - Quality review & scoring",
            "Agent 10: ExportMaster (Haiku) - PDF/DOCX/Google Docs export",
            "Agent 11: CoverDesigner (Sonnet+DALL-E) - AI cover generation",
            "Agent 12: DeliveryManager (Haiku) - Final assembly & delivery",
        ]),
        ("ContentWriter - Topic-by-Topic Flow", [
            "Writes ONE topic at a time with user style selection.",
            "7 Writing Styles: Storytelling, Step-by-Step, Conversational,",
            "  Academic, Motivational, Q&A, Case Study",
            "7 Rewrite Angles: More practical, More emotional, More data-driven,",
            "  Simpler language, More detailed, Different example, Shorter",
            "",
            "Flow: Ask style -> Write topic -> Preview design -> User confirms",
            "  or rewrites with different angle -> Next topic",
        ]),
        ("Financial Model", [
            "Cost per ebook: BDT 173-255 (~$1.57-2.32 USD)",
            "",
            "Pricing Plans:",
            "  Free: BDT 0 - 1 ebook/month, PDF only",
            "  Pro: BDT 999/month - 10 ebooks, all formats, AI covers",
            "  Unlimited: BDT 2,499/month - Unlimited, all features",
            "",
            "Break-even: 32 Pro subscribers",
            "Month 6 profit projection: BDT 107,350",
            "Month 12 profit projection: BDT 399,620",
            "Year 2 profit projection: BDT 1,098,800/month",
            "",
            "LTV:CAC Ratio: 6:1 - 12:1",
            "Gross margin (Pro): 60-75%",
        ]),
        ("Token Management & Auto-Resume", [
            "Smart token management with auto-pause and resume.",
            "Each topic completion = checkpoint saved to DB.",
            "If credits run out: pause, save state, notify user.",
            "On resume: read context, continue from exact stop point.",
            "",
            "Total tokens per ebook: ~157,000",
            "ContentWriter: ~4,000 tokens per topic (critical)",
        ]),
        ("Design & Export Features", [
            "7 Bangla font options for ebook body text",
            "5 header styles, 5 footer styles",
            "Page decorations: quote boxes, tip boxes, stat callouts",
            "Responsive preview: laptop (700px) + mobile (full-width)",
            "Dark mode, Light mode, Sepia mode",
            "Font size slider (14px - 24px)",
            "Export: PDF, DOCX, Google Docs with full styling",
            "",
            "Cover: DALL-E 3 / fal.ai background art",
            "+ Bangla text overlay via Pillow (Python)",
        ]),
        ("Build Timeline (7 Phases, 37 Days)", [
            "Phase 1 (Days 1-4): Foundation setup",
            "Phase 2 (Days 5-9): Agents 1-4 + UI",
            "Phase 3 (Days 10-14): Agents 5-7 + UI",
            "Phase 4 (Days 15-21): ContentWriter (most complex)",
            "Phase 5 (Days 22-26): Export + Design system",
            "Phase 6 (Days 27-32): Covers + SSLCommerz/bKash",
            "Phase 7 (Days 33-37): Polish + Launch",
        ]),
    ]

    for title, lines in sections:
        pdf.add_page()
        pdf.set_font('Helvetica', 'B', 16)
        pdf.set_text_color(30, 58, 138)
        pdf.cell(0, 12, title, 0, 1, 'L')

        pdf.set_draw_color(255, 107, 53)
        pdf.set_line_width(0.3)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(5)

        pdf.set_font('Helvetica', '', 10.5)
        pdf.set_text_color(30, 41, 59)
        for line in lines:
            if line == "":
                pdf.ln(3)
            elif line.startswith("Agent ") or line.startswith("Phase "):
                pdf.set_font('Helvetica', 'B', 10.5)
                pdf.set_text_color(255, 107, 53)
                pdf.cell(0, 7, line, 0, 1)
                pdf.set_font('Helvetica', '', 10.5)
                pdf.set_text_color(30, 41, 59)
            else:
                pdf.cell(0, 7, line, 0, 1)

    pdf_path = os.path.join(OUTPUT_DIR, "OTA_Ebook_Creator_PRD.pdf")
    pdf.output(pdf_path)
    print(f"PDF saved: {pdf_path}")
    return pdf_path


if __name__ == "__main__":
    docx_path = create_prd_docx()
    pdf_path = create_prd_pdf(docx_path)
    print(f"\nDone! Files created:")
    print(f"  DOCX: {docx_path}")
    print(f"  PDF:  {pdf_path}")
